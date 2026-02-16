"""Utilities for rendering Gemini analysis output into a DOCX report."""

from __future__ import annotations

from datetime import datetime
from io import BytesIO
from typing import Any, Dict, Iterable, List

try:
    from docx import Document
    from docx.shared import Pt
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False


def _format_value(value: Any) -> str:
    if value is None:
        return "Not specified"
    if isinstance(value, (list, tuple, set)):
        filtered = [str(item) for item in value if item]
        return ", ".join(filtered) if filtered else "Not specified"
    return str(value)


def _add_key_value_table(document: Document, rows: Iterable[tuple[str, Any]]) -> None:
    rows = [(key, _format_value(value)) for key, value in rows if key]
    if not rows:
        return

    table = document.add_table(rows=len(rows), cols=2)
    table.style = "Light Grid"
    for (key, value), row in zip(rows, table.rows):
        row.cells[0].text = key
        row.cells[1].text = value


def _add_bullet_list(document: Document, items: Iterable[str]) -> None:
    for item in items:
        if not item:
            continue
        paragraph = document.add_paragraph(item, style="List Bullet")
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_after = Pt(3)


def _format_page_label(page: Any) -> str:
    if page in (None, "", "?"):
        return "Page ?"
    return f"Page {page}"


def build_gemini_report(results: Dict[str, Any]) -> BytesIO:
    """Render Gemini AI analysis results into a downloadable DOCX file."""
    if not _DOCX_AVAILABLE:
        raise RuntimeError("python-docx is required for report generation. Install with: pip install python-docx")

    document = Document()
    heading = document.add_heading("Gemini Fire Alarm Detailed Report", level=0)
    heading.alignment = 0

    analysis_ts = results.get("analysis_timestamp")
    if analysis_ts:
        try:
            parsed = datetime.fromisoformat(analysis_ts)
        except ValueError:
            parsed = None
    else:
        parsed = None

    generated_text = parsed.strftime("%B %d, %Y at %I:%M %p") if parsed else "Unknown"
    intro = document.add_paragraph()
    intro.add_run(f"Generated: {generated_text}\n").bold = True
    total_pages = results.get("total_pages")
    if total_pages is not None:
        intro.add_run(f"Total Pages Reviewed: {total_pages}\n")
    if job_id := results.get("job_id"):
        intro.add_run(f"Gemini Job ID: {job_id}")

    # Project Overview
    project_info = results.get("project_info", {}) or {}
    document.add_heading("Project Overview", level=1)
    _add_key_value_table(
        document,
        [
            ("Project Name", project_info.get("project_name")),
            ("Location", project_info.get("location")),
            ("Project Type", project_info.get("project_type")),
            ("Owner / Client", project_info.get("owner")),
            ("Architect", project_info.get("architect")),
            ("Engineer", project_info.get("engineer")),
            ("Project Number", project_info.get("project_number")),
        ],
    )
    if scope := project_info.get("scope_summary"):
        document.add_paragraph(scope)

    # Codes & Standards
    document.add_heading("Fire Alarm Codes & Standards", level=1)
    code_requirements = results.get("code_requirements", {}) or {}
    codes = code_requirements.get("fire_alarm_codes") or code_requirements.get("fire_alarm_standards") or []
    if codes:
        _add_bullet_list(document, codes)
    else:
        document.add_paragraph("No fire alarm-specific codes were extracted.")

    # Fire Alarm Focus Pages
    document.add_heading("Fire Alarm Focus Pages", level=1)
    fa_pages: List[int] = results.get("fire_alarm_pages") or []
    if fa_pages:
        document.add_paragraph(
            "Gemini isolated the following sheets as containing electrical or life-safety fire alarm content:"
        )
        _add_bullet_list(document, [f"{_format_page_label(page)}" for page in fa_pages])
    else:
        document.add_paragraph("No specific fire alarm pages were identified.")

    # Fire Alarm Notes
    document.add_heading("Fire Alarm System Notes", level=1)
    fire_alarm_notes = results.get("fire_alarm_notes") or []
    if fire_alarm_notes:
        for note in fire_alarm_notes:
            if not isinstance(note, dict):
                continue
            page = _format_page_label(note.get("page"))
            note_type = note.get("note_type") or "Note"
            content = note.get("content") or "Details unavailable"
            paragraph = document.add_paragraph(style="List Number")
            paragraph.add_run(f"{page} – {note_type}: ").bold = True
            paragraph.add_run(content)
    else:
        document.add_paragraph("No project-specific fire alarm notes were captured.")

    # Mechanical coordination (duct detectors, dampers)
    document.add_heading("Mechanical Coordination (Fire Alarm Tie-Ins)", level=1)
    mechanical_devices = results.get("mechanical_devices") or {}
    mech_sections = [
        ("Duct Detectors", mechanical_devices.get("duct_detectors", [])),
        ("Fire / Smoke Dampers", mechanical_devices.get("dampers", [])),
        ("HVAC Equipment Over 2000 CFM", mechanical_devices.get("high_airflow_units", [])),
    ]
    added_mechanical = False
    for title, devices in mech_sections:
        document.add_paragraph(title, style="Intense Quote")
        if devices:
            added_mechanical = True
            for device in devices:
                if not isinstance(device, dict):
                    continue
                page = _format_page_label(device.get("page"))
                airflow = device.get("airflow_cfm")
                damper_type = device.get("damper_type")
                fire_alarm_action = device.get("fire_alarm_action")
                requires_dd = device.get("requires_duct_detector")

                parts = [
                    f"{page}",
                    device.get("device_type"),
                    device.get("location") or device.get("equipment_id"),
                    f"Qty: {device.get('quantity')}" if device.get("quantity") else None,
                    f"Airflow: {airflow} CFM" if airflow else None,
                    damper_type,
                    f"Duct detector: {requires_dd}" if requires_dd else None,
                    fire_alarm_action,
                    device.get("specifications"),
                ]
                text = " | ".join(filter(None, map(str, parts)))
                document.add_paragraph(text, style="List Bullet")
        else:
            document.add_paragraph("No devices noted in this category.", style="List Bullet")

    if not added_mechanical:
        document.add_paragraph(
            "Gemini did not identify duct detectors or fire/smoke dampers tied into the fire alarm system."
        )

    # Device placement review
    device_layout = results.get("device_layout_review") or {}
    document.add_heading("Device Placement Review", level=1)
    primary_page = device_layout.get("primary_fa_page") or {}
    unusual = device_layout.get("unusual_placements") or []
    co_detection = device_layout.get("co_detection") or {}

    if primary_page:
        page_label = _format_page_label(primary_page.get("page"))
        reason = primary_page.get("reason") or primary_page.get("note")
        text = f"Most fire alarm devices appear on {page_label}"
        if reason:
            text += f" – {reason}"
        document.add_paragraph(text)

    if unusual:
        document.add_paragraph("Unusual placements or explanations:")
        for item in unusual:
            page = _format_page_label(item.get("page"))
            label = item.get("device_type") or "Device"
            placement = item.get("placement")
            reason = item.get("reason") or item.get("impact")
            text = f"{page} | {label}"
            if placement:
                text += f" – {placement}"
            if reason:
                text += f" ({reason})"
            document.add_paragraph(text, style="List Bullet")

    if co_detection:
        co_needed = co_detection.get("needed")
        co_reason = co_detection.get("reason")
        paragraph = document.add_paragraph()
        paragraph.add_run("CO Detection: ").bold = True
        co_text = co_needed or "Unknown"
        if co_reason:
            co_text += f" – {co_reason}"
        paragraph.add_run(co_text)
    elif not (primary_page or unusual):
        document.add_paragraph("No specific device placement details were captured.")

    # Specifications & manufacturers
    document.add_heading("Fire Alarm Specifications", level=1)
    specifications = results.get("specifications") or {}
    if specifications:
        for key, value in specifications.items():
            if not value or key in {"error"}:
                continue
            label = key.replace("_", " ").title()
            paragraph = document.add_paragraph()
            paragraph.add_run(f"{label}: ").bold = True
            paragraph.add_run(_format_value(value))
    else:
        document.add_paragraph("No additional specification details were captured by Gemini.")

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer

